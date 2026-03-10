"""Ядро объединения DOCX-документов на бесплатном стеке python-docx + docxcompose."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from tempfile import NamedTemporaryFile
from threading import Event
from typing import Callable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docxcompose.composer import Composer

from .logger import get_logger
from .platform_utils import normalize_path


class DocumentMergeError(Exception):
    """Базовая ошибка модуля объединения документов."""


class MergeCancelledError(DocumentMergeError):
    """Ошибка штатной отмены операции пользователем."""


class EmptyFileListError(DocumentMergeError):
    """Ошибка пустого списка файлов для объединения."""


class UnsupportedFormatError(DocumentMergeError):
    """Ошибка неподдерживаемого формата входного файла."""


class FileAccessError(DocumentMergeError):
    """Ошибка доступа к входному или выходному файлу."""


class CorruptedDocumentError(DocumentMergeError):
    """Ошибка чтения поврежденного DOCX-документа."""


class OutputWriteError(DocumentMergeError):
    """Ошибка записи итогового файла."""


@dataclass(slots=True)
class MergeRequest:
    """Параметры задачи объединения документов."""

    source_files: list[Path]
    output_file: Path
    merge_mode: str = "page_break"


@dataclass(slots=True)
class MergeProgress:
    """Данные для отображения прогресса в интерфейсе."""

    current: int
    total: int
    message: str


class DocumentMerger:
    """Объединяет DOCX-файлы с максимально возможным сохранением форматирования бесплатными средствами."""

    SUPPORTED_MERGE_MODES = {"page_break", "section_break", "no_break"}

    def __init__(self) -> None:
        self._logger = get_logger(__name__)
        self._files: list[Path] = []

    @property
    def files(self) -> list[Path]:
        """Возвращает текущий список файлов в порядке объединения."""
        return list(self._files)

    def add_file(self, file_path: str | Path) -> bool:
        """Добавляет DOCX-файл в очередь, если он еще не присутствует."""
        normalized_path = normalize_path(file_path)
        self._validate_single_file_metadata(normalized_path)

        if normalized_path in self._files:
            self._logger.info("Файл уже присутствует в очереди: %s", normalized_path)
            return False

        self._files.append(normalized_path)
        self._logger.info("Файл добавлен в очередь объединения: %s", normalized_path)
        return True

    def remove_file(self, file_path: str | Path) -> bool:
        """Удаляет файл из очереди объединения."""
        normalized_path = normalize_path(file_path)
        try:
            self._files.remove(normalized_path)
        except ValueError:
            self._logger.info("Файл отсутствует в очереди и не был удален: %s", normalized_path)
            return False

        self._logger.info("Файл удален из очереди объединения: %s", normalized_path)
        return True

    def clear_files(self) -> None:
        """Полностью очищает очередь файлов для объединения."""
        self._files.clear()
        self._logger.info("Очередь файлов для объединения очищена.")

    def validate_files(self, files: list[str | Path] | None = None) -> list[Path]:
        """Проверяет список файлов на доступность, формат и читаемость."""
        paths = [normalize_path(path) for path in files] if files is not None else list(self._files)
        if not paths:
            raise EmptyFileListError("Список файлов для объединения пуст.")

        validated_paths: list[Path] = []
        for path in paths:
            self._validate_single_file_metadata(path)
            self._validate_document_readable(path)
            validated_paths.append(path)

        return validated_paths

    def merge_documents(
        self,
        output_file: str | Path,
        merge_mode: str = "page_break",
        files: list[str | Path] | None = None,
        cancel_event: Event | None = None,
        progress_callback: Callable[[MergeProgress], None] | None = None,
    ) -> Path:
        """Объединяет документы в итоговый DOCX-файл."""
        self._validate_merge_mode(merge_mode)

        cancel_event = cancel_event or Event()
        source_files = self.validate_files(files)
        output_path = normalize_path(output_file)
        self._validate_output_path(output_path, source_files)

        self._logger.info(
            "Старт объединения %s файлов в %s с режимом %s.",
            len(source_files),
            output_path,
            merge_mode,
        )

        self._check_cancel(cancel_event)

        try:
            destination_document = Document(str(source_files[0]))
        except Exception as exc:
            raise CorruptedDocumentError(
                f"Не удалось открыть исходный документ: {source_files[0]}"
            ) from exc

        composer = Composer(destination_document)
        self._emit_progress(
            progress_callback,
            current=1,
            total=len(source_files),
            message=f"Загружен файл 1 из {len(source_files)}: {source_files[0].name}",
        )

        for index, source_path in enumerate(source_files[1:], start=2):
            self._check_cancel(cancel_event)
            self._apply_merge_mode_to_destination(destination_document, merge_mode)
            self._append_document(composer, self._load_document(source_path))
            self._emit_progress(
                progress_callback,
                current=index,
                total=len(source_files),
                message=f"Обработан файл {index} из {len(source_files)}: {source_path.name}",
            )

        self._check_cancel(cancel_event)
        saved_path = self._save_document_atomically(composer, output_path)
        self._emit_progress(
            progress_callback,
            current=len(source_files),
            total=len(source_files),
            message="Объединение завершено",
        )
        self._logger.info("Итоговый документ успешно сохранен: %s", saved_path)
        return saved_path

    def merge(
        self,
        request: MergeRequest,
        cancel_event: Event,
        progress_callback: Callable[[MergeProgress], None] | None = None,
    ) -> Path:
        """Совместимый адаптер для текущего GUI и worker."""
        return self.merge_documents(
            output_file=request.output_file,
            merge_mode=request.merge_mode,
            files=request.source_files,
            cancel_event=cancel_event,
            progress_callback=progress_callback,
        )

    def _validate_single_file_metadata(self, file_path: Path) -> None:
        """Проверяет базовую доступность файла до открытия."""
        if file_path.suffix.lower() != ".docx":
            raise UnsupportedFormatError(f"Поддерживаются только DOCX-файлы: {file_path}")

        if not file_path.exists():
            raise FileAccessError(f"Файл не найден: {file_path}")

        if not file_path.is_file():
            raise FileAccessError(f"Ожидается файл, а не папка: {file_path}")

        try:
            with file_path.open("rb"):
                pass
        except OSError as exc:
            raise FileAccessError(f"Нет доступа к файлу: {file_path}") from exc

    def _validate_document_readable(self, file_path: Path) -> None:
        """Проверяет, что DOCX-файл действительно открывается движком объединения."""
        try:
            Document(str(file_path))
        except Exception as exc:
            raise CorruptedDocumentError(
                f"Файл поврежден или не может быть прочитан как DOCX: {file_path}"
            ) from exc

    def _load_document(self, file_path: Path) -> Document:
        """Открывает документ и преобразует ошибки чтения в доменные исключения."""
        try:
            return Document(str(file_path))
        except Exception as exc:
            raise CorruptedDocumentError(
                f"Не удалось прочитать DOCX-документ: {file_path}"
            ) from exc

    @staticmethod
    def _apply_merge_mode_to_destination(document: Document, merge_mode: str) -> None:
        """Задает разделитель между документами на стороне результирующего файла.

        Важно: мы не пересохраняем присоединяемый документ через python-docx, чтобы не
        ломать сложную верстку, изображения, колонтитулы и другие элементы OOXML.
        """
        if merge_mode == "page_break":
            document.add_page_break()
            return

        if merge_mode == "section_break":
            document.add_section(WD_SECTION_START.NEW_PAGE)
            return

        if merge_mode == "no_break":
            return

    @staticmethod
    def _append_document(composer: Composer, source: Document) -> None:
        """Присоединяет документ, сохраняя структуру бесплатными средствами."""
        composer.append(source)

    def _save_document_atomically(self, composer: Composer, output_path: Path) -> Path:
        """Сохраняет документ через временный файл и затем перемещает его в целевой путь."""
        output_path.parent.mkdir(parents=True, exist_ok=True)

        try:
            with NamedTemporaryFile(
                suffix=".docx",
                prefix="docx_merger_",
                dir=output_path.parent,
                delete=False,
            ) as temp_file:
                temp_path = Path(temp_file.name)
        except OSError as exc:
            raise OutputWriteError(
                f"Не удалось создать временный файл для сохранения результата: {output_path.parent}"
            ) from exc

        try:
            composer.save(str(temp_path))
            temp_path.replace(output_path)
            return output_path
        except OSError as exc:
            if temp_path.exists():
                temp_path.unlink(missing_ok=True)
            raise OutputWriteError(f"Не удалось записать итоговый файл: {output_path}") from exc
        except Exception as exc:
            if temp_path.exists():
                temp_path.unlink(missing_ok=True)
            raise OutputWriteError(f"Ошибка сохранения итогового документа: {output_path}") from exc

    def _validate_merge_mode(self, merge_mode: str) -> None:
        """Проверяет, что выбранный режим объединения поддерживается."""
        if merge_mode not in self.SUPPORTED_MERGE_MODES:
            supported_modes = ", ".join(sorted(self.SUPPORTED_MERGE_MODES))
            raise ValueError(
                f"Неподдерживаемый режим объединения: {merge_mode}. Доступно: {supported_modes}"
            )

    @staticmethod
    def _validate_output_path(output_path: Path, source_files: list[Path]) -> None:
        """Проверяет, что итоговый файл не совпадает ни с одним из входных документов."""
        if output_path in source_files:
            raise OutputWriteError(
                "Итоговый файл не должен совпадать ни с одним из исходных документов."
            )

    @staticmethod
    def _emit_progress(
        callback: Callable[[MergeProgress], None] | None,
        current: int,
        total: int,
        message: str,
    ) -> None:
        """Отправляет состояние прогресса наружу, если подписчик задан."""
        if callback is not None:
            callback(MergeProgress(current=current, total=total, message=message))

    @staticmethod
    def _check_cancel(cancel_event: Event) -> None:
        """Проверяет запрос на отмену и прерывает операцию на безопасной точке."""
        if cancel_event.is_set():
            raise MergeCancelledError("Операция отменена пользователем.")


# Совместимость с уже существующим GUI и worker.
DocxMerger = DocumentMerger
